# -*- coding: utf-8 -*-
"""TSMC Q1 2026 post-earnings view note (DOCX)."""
import os
os.makedirs("./out", exist_ok=True)
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CJK="Microsoft JhengHei"; LAT="Times New Roman"
NAVY=RGBColor(0x1F,0x4E,0x79); WHITE=RGBColor(0xFF,0xFF,0xFF); GREY=RGBColor(0x55,0x55,0x55)
GREEN=RGBColor(0x1E,0x7A,0x46); RED=RGBColor(0xB0,0x2A,0x2A); BLACK=RGBColor(0x11,0x11,0x11)

doc=Document()
st=doc.styles["Normal"]; st.font.name=LAT; st.font.size=Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"),CJK)

def setcjk(run): run._element.rPr.rFonts.set(qn("w:eastAsia"),CJK)
def para(text,size=10.5,bold=False,color=BLACK,align=None,after=4,before=0,italic=False):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before)
    if align: p.alignment=align
    r=p.add_run(text); r.font.size=Pt(size); r.font.bold=bold; r.font.color.rgb=color; r.font.name=LAT; r.italic=italic
    setcjk(r); return p
def shade(cell,color):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd"); sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),color); tcPr.append(sh)
def setcell(cell,text,size=9,bold=False,color=BLACK,fill=None,align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text=""; p=cell.paragraphs[0]; p.alignment=align; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(text); r.font.size=Pt(size); r.font.bold=bold; r.font.color.rgb=color; r.font.name=LAT; setcjk(r)
    if fill: shade(cell,fill)

# ---- Title banner ----
p=para("台積電 TSMC (2330.TW / TSM) — 財報後觀點筆記",15,True,NAVY,after=1)
para("2026 年第一季 (Q1 2026) | 發布 2026-04-16 | 季末 2026-03-31 | 觀點筆記 (Post-Earnings Note)",9,False,GREY,after=2)
para("實際 vs. 市場預期 vs. 前次指引｜評等與目標價未由本筆記設定；數字未經獨立查核請見來源",9,False,GREY,italic=True,after=8)

# ---- 1. 摘要 ----
para("一、結論摘要 (Beat-and-Raise，但股價「利多出盡」)",12.5,True,NAVY,before=4,after=4)
para("台積電 Q1'26 繳出乾淨的「超預期 + 上修」(beat-and-raise)：營收 US$35.90bn(+40.6% YoY)高於前次指引上緣;"
     "毛利率 66.2% 大幅優於 63–65% 指引(超出上緣約 120bps);稀釋 EPS NT$22.08(+58% YoY)勝市場預期約 6%。"
     "Q2'26 指引中點 US$39.6bn 亦高於市場預期約 4%,全年維持「US$ 營收成長逾 30%」,並上調長期毛利率底線至「56% 以上」。", after=4)
para("然而 ADR 於財報後一個交易日仍下跌約 3%——關鍵不在獲利,而在資本支出:2026 capex 指引偏向 US$52–56bn 高端(創高),"
     "引發短期自由現金流/資本密集度疑慮;加上股價於財報前已較一年前漲約 1 倍、逼近 52 週高點,形成典型「利多出盡」格局。", after=6)

# ---- 2. Actual vs Cons vs Prior table ----
para("二、實際 vs. 市場預期 vs. 前次指引",12.5,True,NAVY,after=4)
t=doc.add_table(rows=1,cols=6); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
hd=["指標","前次指引 (Jan'26)","市場預期","實際","實際 vs 指引","實際 vs 預期"]
for i,h in enumerate(hd): setcell(t.rows[0].cells[i],h,9,True,WHITE,"1F4E79",WD_ALIGN_PARAGRAPH.CENTER)
rows=[
 ("營收 (US$bn)","34.6–35.8","35.5","35.90","高於上緣","+1.1%"),
 ("毛利率","63–65%","~64%","66.2%","+~120bps","+~220bps"),
 ("營業利益率","—","—","58.1%","高於指引","—"),
 ("EPS (NT$)","(無指引)","20.88","22.08","—","+5.7%"),
 ("EPS/ADR (US$)","—","3.26","3.49","—","+7.1%"),
 ("淨利 (NT$bn)","—","—","572.48","—","+58% YoY"),
]
for row in rows:
    cells=t.add_row().cells
    for i,v in enumerate(row):
        fill="E2EFDA" if (i>=4 and (v.startswith("+") or "高於" in v)) else None
        setcell(cells[i],v,9,i==0,BLACK,fill,WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER)
para("資料:TSMC 6-K & Q1'26 法說會(2026-04-16);前次指引出自 Q4'25 財報(2026-01-15);consensus 為第三方彙整,發布前請以 FactSet/Bloomberg 校正。",8,False,GREY,before=2,after=6,italic=True)

para("Q2'26 指引同樣優於預期:營收指引 US$39.0–40.2bn(中點 39.6)vs 市場約 38.1 → 高約 4%;毛利率指引 65.5–67.5%。",10.5,False,BLACK,after=6)

# ---- 3. 法說會重點 ----
para("三、法說會重點 (Earnings Call Highlights)",12.5,True,NAVY,after=4)
bullets=[
 ("AI/HPC 需求","CEO 魏哲家:AI 需求「持續強勁」,並從生成式 AI 轉向「代理式 AI(agentic)與 command-and-action」,運算需求更大;駁斥泡沫論但強調紀律投資。HPC 佔營收 61%。"),
 ("N2 (2nm) 進度","Q4'25 已進入量產且良率良好,於新竹與高雄分階段擴產,約 2026 年 3 月起貢獻營收;尚未單獨拆分佔比。"),
 ("毛利率展望","Q1 受惠高稼動率、成本改善、有利匯率;但明示 H2'26 起 N2 量產 + 海外廠將稀釋毛利率,全年稀釋約 2–3%(海外早期 2–3%、後期擴大至 3–4%)。N3 預計 H2'26 達公司平均毛利率。"),
 ("資本支出","2026 capex 指引偏 US$52–56bn 高端(創高),約 70–80% 投先進製程、其餘為先進封裝(含 CoWoS)與特殊製程——市場負面反應主因。"),
 ("海外布局","亞利桑那二廠(3nm)2027 H2 量產、三廠興建中;日本二廠 2028;台南新 3nm 廠 2027 H1。"),
 ("長期目標上修","長期毛利率底線上調至「56% 以上 through the cycle」;長期 AI 加速器營收 CAGR 上調至 ~54–56%(至 2029)。"),
 ("成本/地緣","示警零組件漲價(消費性敏感);中東地緣推升製程用化學品與氣體價格。Q2 指引匯率假設 USD/TWD = 31.7。"),
]
for h,b in bullets:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); p.paragraph_format.left_indent=Pt(10)
    r=p.add_run("▍"+h+"  "); r.font.bold=True; r.font.size=Pt(10.5); r.font.color.rgb=NAVY; r.font.name=LAT; setcjk(r)
    r2=p.add_run(b); r2.font.size=Pt(10.5); r2.font.color.rgb=BLACK; r2.font.name=LAT; setcjk(r2)

# ---- 4. 模型更新 ----
para("四、模型更新摘要 (Estimate Refresh)",12.5,True,NAVY,before=4,after=4)
para("已新建 TSMC 季度估計模型(工作區先前無台積電模型),以 Q1'26 實際 + Q2'26 指引為錨,H2'26 為標記清楚的估計:", after=3)
t2=doc.add_table(rows=1,cols=6); t2.alignment=WD_TABLE_ALIGNMENT.CENTER; t2.style="Table Grid"
h2=["US$ / NT$","Q1'26A","Q2'26E","Q3'26E","Q4'26E","FY2026E"]
for i,h in enumerate(h2): setcell(t2.rows[0].cells[i],h,9,True,WHITE,"1F4E79",WD_ALIGN_PARAGRAPH.CENTER)
mrows=[
 ("營收 (US$bn)","35.90","39.60","42.77","45.33","163.6"),
 ("YoY","+40.6%","~+32%","—","—","+31.9%"),
 ("毛利率","66.2%","66.5%","65.5%","65.0%","65.8%"),
 ("EPS (NT$)","22.08","23.97","25.62","26.88","98.6"),
]
for row in mrows:
    cells=t2.add_row().cells
    for i,v in enumerate(row):
        setcell(cells[i],v,9,i==0,BLACK,None,WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER)
para("假設:Q3/Q4 營收 +8%/+6% QoQ;H2 毛利率反映 N2+海外稀釋;匯率 31.7;股數 25.93bn(1 ADR=5 股)。"
     "FY2025 基數 ~US$124bn 為近似值,須以官方數字校正。檔案:out/TSMC_Q1-2026_Model.xlsx(3 分頁,含變異表與來源)。",8,False,GREY,before=2,after=6,italic=True)

# ---- 5. 觀點與風險 ----
para("五、觀點與風險 (View & Risks)",12.5,True,NAVY,after=4)
para("多方:AI/HPC 需求廣化(生成式 → 代理式),先進節點佔晶圓營收 74%,毛利率結構性走強(長期底線上調至 56%+),"
     "N3 H2'26 轉正貢獻,定價權與規模優勢延續;Q2 指引續優。", after=3)
para("空方/觀察:① capex 偏高端壓抑近期 FCF(本季股價負反應主因);② N2 量產 H2 起稀釋全年毛利率 ~2–3%;"
     "③ 海外廠稀釋 2–3%→後期 3–4%;④ 強勢台幣對 US$ 報表/ADR 為結構性逆風(本季雖為順風);⑤ 客戶庫存、手機復甦、地緣與關稅不確定。", after=4)
para("淨評:基本面強勁且長期目標上修,屬高品質持有;但股價已反映高度樂觀,短期催化偏向 capex/FCF 與 H2 毛利率稀釋節奏。"
     "建議把觀察重心放在『N2 放量曲線 vs. 稀釋幅度』與『資本密集度何時見頂』。",10.5,True,BLACK,after=8)

# ---- Sources ----
para("來源 (Sources)",11.5,True,NAVY,after=3)
src=[
 "TSMC 6-K 新聞稿「TSMC Reports First Quarter EPS of NT$22.08」(2026-04-16) — pr.tsmc.com/english/news/3297",
 "TSMC Q1 2026 法說會 transcript — investing.com(Q1'26 earnings call transcript)",
 "TSMC Q4'25 結果 / Q1'26 前次指引(2026-01-15) — investor.tsmc.com/english/quarterly-results/2025/q4",
 "Manufacturing Dive、DataCenterDynamics(結果/指引/capex/評論)",
 "Tickeron、TipRanks(consensus ~US$35.5bn、EPS US$3.26/NT$20.88、股價反應)",
 "BofA 目標價 NT$2,360→2,530(維持 Buy)— 第三方彙整,須查核",
]
for s in src:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2); p.paragraph_format.left_indent=Pt(10)
    r=p.add_run("• "+s); r.font.size=Pt(8.5); r.font.color.rgb=GREY; r.font.name=LAT; setcjk(r)
para("免責:本筆記彙整自公開資訊與第三方來源,僅供內部研究參考,不構成投資建議。consensus、股價反應與目標價發布前請以 FactSet/Bloomberg 校正;FY2025 基數與地區別營收待官方數字補正。",8,False,GREY,before=4,italic=True)

doc.save("./out/TSMC_Q1_2026_Earnings_Note.docx")
print("SAVED ./out/TSMC_Q1_2026_Earnings_Note.docx  paragraphs:", len(doc.paragraphs))
