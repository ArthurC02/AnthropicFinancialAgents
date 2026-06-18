# -*- coding: utf-8 -*-
"""TSMC Q1 2026 full earnings update report (8-12pp DOCX with charts)."""
import os
os.makedirs("./out", exist_ok=True)
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CJK="Microsoft JhengHei"; LAT="Times New Roman"
NAVY=RGBColor(0x1F,0x4E,0x79); WHITE=RGBColor(0xFF,0xFF,0xFF); GREY=RGBColor(0x55,0x55,0x55)
GREEN=RGBColor(0x1E,0x7A,0x46); RED=RGBColor(0xB0,0x2A,0x2A); BLACK=RGBColor(0x11,0x11,0x11); AMBER=RGBColor(0x9A,0x66,0x00)
CH="./out/charts"

doc=Document()
sec=doc.sections[0]
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(sec,m,Inches(0.8))
st=doc.styles["Normal"]; st.font.name=LAT; st.font.size=Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"),CJK)

def cjk(run): run._element.rPr.rFonts.set(qn("w:eastAsia"),CJK)
def P(text,size=10.5,bold=False,color=BLACK,align=None,after=4,before=0,italic=False):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before)
    if align is not None: p.alignment=align
    r=p.add_run(text); r.font.size=Pt(size); r.font.bold=bold; r.font.color.rgb=color; r.font.name=LAT; r.italic=italic; cjk(r)
    return p
def H(text,size=13,before=8,after=4):
    P(text,size,True,NAVY,before=before,after=after)
def bullet(head,body,color=NAVY):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); p.paragraph_format.left_indent=Pt(12)
    r=p.add_run("▍"+head+"  "); r.font.bold=True; r.font.size=Pt(10.5); r.font.color.rgb=color; r.font.name=LAT; cjk(r)
    r2=p.add_run(body); r2.font.size=Pt(10.5); r2.font.color.rgb=BLACK; r2.font.name=LAT; cjk(r2)
def shade(cell,c):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd"); sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),c); tcPr.append(sh)
def cell(c,text,size=9,bold=False,color=BLACK,fill=None,align=WD_ALIGN_PARAGRAPH.LEFT):
    c.text=""; p=c.paragraphs[0]; p.alignment=align; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(text); r.font.size=Pt(size); r.font.bold=bold; r.font.color.rgb=color; r.font.name=LAT; cjk(r)
    if fill: shade(c,fill)
def chart(name,width=6.6,caption=None):
    doc.add_picture(f"{CH}/{name}.png",width=Inches(width))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    if caption: P(caption,8,False,GREY,align=WD_ALIGN_PARAGRAPH.CENTER,after=6,italic=True)

# ============ PAGE 1 — SUMMARY ============
P("股票研究 · 財報更新 EQUITY RESEARCH — EARNINGS UPDATE",9,True,NAVY,after=1)
P("台積電 TSMC（2330.TW / NYSE: TSM）",18,True,NAVY,after=1)
P("2026 年第一季財報更新｜發布 2026-04-16｜季末 2026-03-31",10,False,GREY,after=6)

# rating box (table)
tb=doc.add_table(rows=1,cols=4); tb.alignment=WD_TABLE_ALIGNMENT.CENTER; tb.style="Table Grid"
box=[("本報告觀點","中性偏多 (Hold-plus)"),("外部目標價參考","BofA NT$2,530 (Buy)"),
     ("Q1'26 EPS","NT$22.08 (+58% YoY)"),("結果","Beat-and-Raise")]
for i,(k,v) in enumerate(box):
    c=tb.rows[0].cells[i]; c.text=""
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(k+"\n"); r.font.size=Pt(8); r.font.color.rgb=GREY; r.font.name=LAT; cjk(r)
    r2=p.add_run(v); r2.font.size=Pt(9.5); r2.font.bold=True; r2.font.color.rgb=NAVY; r2.font.name=LAT; cjk(r2)
    shade(c,"EAF1F8")
P("",4,after=2)

H("投資要點 (Key Takeaways)",13,before=4)
bullet("乾淨超預期","營收 US$35.90bn(+40.6% YoY)高於前次指引上緣;毛利率 66.2% 超 63–65% 指引上緣約 120bps;EPS NT$22.08 勝市場約 6%。")
bullet("指引續優","Q2'26 營收指引中點 US$39.6bn 高於市場約 4%;全年維持「US$ 營收 >30%」,長期毛利率底線上調至 56%+。")
bullet("但股價利多出盡","ADR 財報後一日仍跌約 3%——主因 2026 capex 指引偏 US$52–56bn 高端(創高)引發 FCF 疑慮,股價又已逼近 52 週高點。")
bullet("觀察重心","N2 放量曲線 vs. 毛利率稀釋(H2 起 N2+海外廠全年稀釋約 2–3%);資本密集度何時見頂。")
P("資料:TSMC 6-K & Q1'26 法說會(2026-04-16);前次指引出自 Q4'25 財報(2026-01-15);consensus 為第三方彙整,發布前請以 FactSet/Bloomberg 校正。",8,False,GREY,before=2,italic=True)

chart("01_revenue",6.6,"圖1:季度營收(US$bn)— Q1'26 實際、Q2'26 指引、H2'26 估計")

# ============ DETAILED RESULTS ============
doc.add_page_break()
H("一、Q1'26 業績詳析:全面優於預期",13)
P("台積電 Q1'26 在淡季繳出全面超預期成績。營收 US$35.90bn,季增 6.4%、年增 40.6%(以新台幣計 NT$1,134.10bn,年增 35.1%);"
  "新台幣與美元成長率差異來自匯率。獲利端更亮眼:毛利率 66.2%(季增 390bps、年增 740bps),營業利益率 58.1%,淨利率 50.5%,"
  "淨利 NT$572.48bn(+58.3% YoY),稀釋 EPS NT$22.08、每 ADR US$3.49。", after=4)
P("相對市場預期:營收勝約 1.1%、EPS 勝約 5.7%、毛利率勝指引上緣約 120bps。毛利率驅動為高稼動率、成本改善、有利匯率與技術組合。", after=4)
chart("06_beatmiss",6.4,"圖2:Q1'26 實際 vs. 市場預期 vs. 前次指引(營收與 EPS)")
chart("03_eps",6.4,"圖3:季度 EPS(NT$)— Q1'26 NT$22.08,年增 58%")

# Summary table 1: results
H("表1:Q1'26 損益摘要",11,before=6,after=3)
t=doc.add_table(rows=1,cols=4); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
for i,h in enumerate(["指標","Q1'26 實際","YoY","QoQ"]): cell(t.rows[0].cells[i],h,9,True,WHITE,"1F4E79",WD_ALIGN_PARAGRAPH.CENTER)
res=[("營收 (US$bn)","35.90","+40.6%","+6.4%"),("營收 (NT$bn)","1,134.10","+35.1%","+8.4%"),
 ("毛利率","66.2%","+740bps","+390bps"),("營業利益率","58.1%","—","—"),
 ("淨利 (NT$bn)","572.48","+58.3%","+13.2%"),("EPS (NT$)","22.08","+58.3%","—"),("EPS/ADR (US$)","3.49","—","—")]
for row in res:
    cs=t.add_row().cells
    for i,v in enumerate(row): cell(cs[i],v,9,i==0,BLACK,None,WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER)
P("資料:TSMC Q1'26 6-K 新聞稿(2026-04-16)。",8,False,GREY,before=2,italic=True)

# ============ REVENUE MIX ============
doc.add_page_break()
H("二、營收組合:AI/HPC 主導,先進製程佔 74%",13)
P("平台別,HPC 佔營收 61%、智慧型手機 26%、IoT 6%、車用 4%、DCE 1%、其他 2%——HPC 佔比續創高,反映 AI/資料中心需求主導的結構轉變。"
  "技術節點,3nm 佔晶圓營收 25%、5nm 36%、7nm 13%,先進製程(≤7nm)合計達 74%。N2(2nm)Q4'25 量產、3 月起開始貢獻,尚未單獨拆分。", after=4)
chart("04_platform",6.4,"圖4:Q1'26 平台別營收 — HPC 61% 主導")
chart("05_node",6.4,"圖5:Q1'26 技術節點別 — 先進製程(≤7nm)佔 74%")

# ============ MARGINS ============
doc.add_page_break()
H("三、毛利率與獲利能力:結構走強,但 H2 面臨稀釋",13)
P("毛利率 66.2% 大幅優於指引,連續第八季成長動能延續,獲益於稼動率、成本改善與匯率順風。惟管理層明示下半年起 N2 量產與海外廠擴張將稀釋毛利率,"
  "全年稀釋約 2–3%(海外早期 2–3%、後期擴大至 3–4%);N3 預計 H2'26 達公司平均毛利率。我們因此在模型中對 H2 毛利率採保守假設(Q3 65.5%、Q4 65.0%)。", after=4)
chart("02_grossmargin",6.4,"圖6:季度毛利率 — Q1'26 66.2% 超指引上緣約 120bps")
chart("10_gmwalk",6.4,"圖7:FY26E 毛利率走勢(示意)— H2 N2 與海外廠稀釋")

# ============ GUIDANCE ============
doc.add_page_break()
H("四、指引與展望:Q2 續優,全年 >30%,資本支出創高",13)
P("Q2'26 指引:營收 US$39.0–40.2bn(中點季增約 10%、年增約 32%),高於市場預期約 4%;毛利率 65.5–67.5%、營業利益率 56.5–58.5%;匯率假設 USD/TWD 31.7。"
  "全年:維持「US$ 營收成長逾 30%」(AI 驅動);2026 資本支出偏 US$52–56bn 高端(創高),約 70–80% 投先進製程,其餘為先進封裝(含 CoWoS)與特殊製程;"
  "長期 AI 加速器營收 CAGR 上調至 ~54–56%(至 2029)。", after=4)
chart("07_q2guide",6.0,"圖8:Q2'26 營收指引較市場預期高約 4%")
chart("08_fyrevenue",6.4,"圖9:全年營收與成長 — FY26E '>30%'(AI 驅動)")
chart("09_capex",6.4,"圖10:2026 資本支出 US$52–56bn 高端 — 本季股價負反應主因")

# ============ THESIS ============
doc.add_page_break()
H("五、投資論點更新 (Thesis)",13)
P("多方:AI/HPC 需求廣化(生成式 → 代理式 agentic,運算需求更大),先進製程佔晶圓營收 74%,毛利率結構性走強且長期底線上調至 56%+,"
  "N3 H2'26 轉正貢獻,定價權與規模優勢延續,Q2 指引續優。", after=3)
P("空方/觀察:① capex 偏高端壓抑近期 FCF(本季股價負反應主因);② N2 量產 H2 起全年稀釋毛利率約 2–3%;③ 海外廠稀釋 2–3%→後期 3–4%;"
  "④ 強勢台幣對美元報表/ADR 為結構性逆風(本季雖為順風);⑤ 客戶庫存、手機復甦、地緣與關稅不確定。", after=4)
P("淨評:基本面強勁且長期目標上修,屬高品質持有(中性偏多)。短期催化偏向 capex/FCF 與 H2 毛利率稀釋節奏;"
  "建議把觀察重心放在『N2 放量曲線 vs. 稀釋幅度』與『資本密集度何時見頂』。",10.5,True,BLACK,after=6)

# ============ ESTIMATES ============
H("六、預估摘要與估值",13,before=4)
P("已新建 TSMC 季度估計模型(工作區先前無台積電模型),以 Q1'26 實際 + Q2'26 指引為錨,H2'26 為估計:", after=3)
t2=doc.add_table(rows=1,cols=6); t2.alignment=WD_TABLE_ALIGNMENT.CENTER; t2.style="Table Grid"
for i,h in enumerate(["項目","Q1'26A","Q2'26E","Q3'26E","Q4'26E","FY2026E"]): cell(t2.rows[0].cells[i],h,9,True,WHITE,"1F4E79",WD_ALIGN_PARAGRAPH.CENTER)
mrows=[("營收 (US$bn)","35.90","39.60","42.77","45.33","163.6"),("YoY","+40.6%","~+32%","—","—","+31.9%"),
 ("毛利率","66.2%","66.5%","65.5%","65.0%","65.8%"),("營業利益率","58.1%","57.5%","56.8%","56.3%","57.0%"),
 ("EPS (NT$)","22.08","23.97","25.62","26.88","98.6")]
for row in mrows:
    cs=t2.add_row().cells
    for i,v in enumerate(row): cell(cs[i],v,9,i==0,BLACK,None,WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER)
P("假設:Q3/Q4 營收 +8%/+6% QoQ;H2 毛利率反映 N2+海外稀釋;匯率 31.7;股數 25.93bn(1 ADR=5 股);FY2025 基數 ~US$124bn 為近似值。"
  "完整估值另以 DCF + 同業 comps 交付(見隨附 Excel)。",8,False,GREY,before=2,italic=True)

# ============ RISKS + SOURCES ============
# ============ VALUATION APPENDIX ============
doc.add_page_break()
H("附錄:估值 (Valuation Appendix — DCF + Comps)",13)
P("我們以美元 DCF(承接季度模型 FY26E 營收 US$163.6bn,顯性期 FY27–31,期中折現,淨現金回加)估值。WACC 9.95%"
  "(Beta 1.15、無風險利率 4.3%、ERP 5%,淨現金 ~US$65bn 使資金成本接近權益成本);終值採永續成長法。", after=4)
P("基準情境隱含每 ADR 約 US$340、折合 NT$2,157/股,相對市場 ~US$380 約 −10%——即「大致合理、市場給約一成的品質溢價」。"
  "三情境區間 US$221(熊)~ US$463(牛),涵蓋市場價與 BofA 目標價 NT$2,530(≈US$399/ADR)。", after=4)
# scenario table
H("表3:DCF 三情境(每 ADR / 每股)",11,before=4,after=3)
t3=doc.add_table(rows=1,cols=5); t3.alignment=WD_TABLE_ALIGNMENT.CENTER; t3.style="Table Grid"
for i,h in enumerate(["情境","每 ADR (US$)","每股 (NT$)","vs 市場 ~$380","TV 占 EV"]):
    cell(t3.rows[0].cells[i],h,9,True,WHITE,"1F4E79",WD_ALIGN_PARAGRAPH.CENTER)
srows=[("熊 Bear","220.67","1,399","−41.9%","73%"),("基 Base","340.19","2,157","−10.5%","78%"),("牛 Bull","462.70","2,934","+21.8%","80%")]
for row in srows:
    cs=t3.add_row().cells
    for i,v in enumerate(row):
        fl="EAF1F8" if row[0].startswith("基") else None
        cell(cs[i],v,9,row[0].startswith("基"),BLACK,fl,WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER)
P("情境驅動:營收成長 / 營業利益率 / Capex 強度 / 終值成長 g;WACC 固定(其變化見工作簿 WACC×g 敏感度表)。",8,False,GREY,before=2,italic=True)
chart("11_football",6.2,"圖11:DCF 三情境隱含每 ADR vs. 市場價與 BofA 目標價")
# comps table
H("表4:同業比較(概略)",11,before=4,after=3)
t4=doc.add_table(rows=1,cols=5); t4.alignment=WD_TABLE_ALIGNMENT.CENTER; t4.style="Table Grid"
for i,h in enumerate(["公司","市值(US$bn)","毛利率","Fwd P/E","EV/EBITDA"]):
    cell(t4.rows[0].cells[i],h,9,True,WHITE,"1F4E79",WD_ALIGN_PARAGRAPH.CENTER)
crows=[("台積電 TSMC","~985","66%","~20x","~12x"),("Samsung","~430","~38%","~13x","~5x"),
 ("UMC 聯電","~18","~32%","~13x","~5x"),("GlobalFoundries","~24","~25%","~18x","~8x"),
 ("SMIC 中芯","~75","~22%","~35x","~9x"),("Intel","~120","~32%","~25x","~8x")]
for row in crows:
    cs=t4.add_row().cells
    for i,v in enumerate(row):
        fl="EAF1F8" if row[0].startswith("台積") else None
        cell(cs[i],v,9,row[0].startswith("台積"),BLACK,fl,WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER)
P("TSMC 毛利率 66% 遠高於同業(22–38%),~20x Fwd P/E / ~12x EV/EBITDA 的溢價由獲利品質與 AI 領導地位支撐;純代工可比者稀少。",8,False,GREY,before=2,italic=True)
P("估值小結:DCF 與 comps 一致指向「高品質、合理偏小幅溢價」。與 Vertiv(同題材但 DCF 隱含約 −68%)相比,TSMC 的折價/溢價溫和,"
  "差異來自低 beta + 淨現金 + 真實 FCF;惟 capex ~33% 的高資本密集度是 FCF 主要拖累,也是本季股價負反應根源。完整模型見 out/TSMC_DCF_Valuation.xlsx(含情境選擇器與敏感度表)。",
  10.5,True,BLACK,before=3,after=6)

H("七、風險",12,before=6)
for h,b in [("資本密集度","capex 偏高端壓抑近期 FCF"),("毛利率稀釋","N2 與海外廠 H2 起全年稀釋約 2–3%"),
            ("匯率","強勢台幣對美元報表/ADR 結構性逆風"),("需求","AI 需求若降溫 + 過度投資的過剩風險"),
            ("地緣/關稅","美中貿易、製程用化學品/氣體漲價")]:
    bullet(h,b,color=RED)

H("來源 (Sources & References)",11,before=6)
for s in ["TSMC 6-K 新聞稿「First Quarter EPS of NT$22.08」(2026-04-16) — pr.tsmc.com/english/news/3297",
          "TSMC Q1'26 法說會 transcript — investing.com",
          "TSMC Q4'25 結果 / Q1'26 前次指引(2026-01-15) — investor.tsmc.com",
          "Manufacturing Dive、DataCenterDynamics、Tickeron、TipRanks(consensus 與股價反應)",
          "BofA 目標價 NT$2,360→2,530(維持 Buy)— 第三方彙整,須查核"]:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2); p.paragraph_format.left_indent=Pt(10)
    r=p.add_run("• "+s); r.font.size=Pt(8.5); r.font.color.rgb=GREY; r.font.name=LAT; cjk(r)
P("免責:本報告彙整自公開資訊與第三方來源,僅供內部研究參考,不構成投資建議。圖表中 2025 各季與資本支出歷史為近似估計(已標註);"
  "consensus、股價反應、目標價、FY2025 基數與地區別營收待官方數字補正。",8,False,GREY,before=4,italic=True)

doc.save("./out/TSMC_Q1_2026_Earnings_Update.docx")
print("SAVED ./out/TSMC_Q1_2026_Earnings_Update.docx  paragraphs:",len(doc.paragraphs))
